from flask import Flask, request, render_template, redirect, session, jsonify,send_file
from mailer import Mailer
import random
import hashlib
import sqlite3
from datetime import date, datetime
import html
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL


daych = datetime.now().date()
date_  = date.today()

app = Flask(__name__)
app.secret_key = '05026528418405108184010841'




def encrippt(prs):

    p1e = hashlib.md5(prs.encode()).hexdigest()
    p1d1 = p1e.translate({ord('b'): None})
    p1d2 = p1d1.translate({ord('8'): None})
    p1 = p1d2[0:9]


    p2e = hashlib.sha1(prs.encode()).hexdigest()
    p2d1 = p2e.translate({ord('e'): None})
    p2d2 = p2d1.translate({ord('6'): None})
    p2 = p2d2[0:6]


    p3e = hashlib.sha256(prs.encode()).hexdigest()
    p3d1 = p3e.translate({ord('1'): None})
    p3d2 = p3d1.translate({ord('l'): None})
    p3 = p3d2[0:10]


    return str(p1)+str(p2)+str(p3)




@app.route('/home')
def homes():

    nameee = session.get('Name')
    Emailll = session.get('Email')
    ss = session.get('Email')

    db = sqlite3.connect("Totringdb.db")
    mydb = db.cursor()

    mydb.execute(f"SELECT * FROM attendees")
    allUser = mydb.fetchall()

    return render_template('home.html',name=nameee,Email=Emailll,ss=ss, allatt=allUser)



@app.route('/adminCCA', methods =["GET", "POST"] )
def homeg():
    if session.get('Email') != 'admin email':
        return redirect('/home')





    nameee = session.get('Name')
    Emailll = session.get('Email')
    ss = session.get('Email')
    db = sqlite3.connect("Totringdb.db")
    mydb = db.cursor()




    # all user accepted
    mydb.execute(f"SELECT * FROM usersz WHERE AdminAccept='True'")
    allUser = mydb.fetchall()


    # Order Accepted
    mydb.execute(f"SELECT * FROM Orders WHERE OrderStatus='True'")
    OrderAccepted = mydb.fetchall()






    # whiting to accept user
    mydb.execute(f"SELECT * FROM usersz WHERE AdminAccept='False'")
    whitingAccept = mydb.fetchall()


    # Order whiting to accept
    mydb.execute(f"SELECT * FROM Orders WHERE OrderStatus='False'")
    whitingOrder = mydb.fetchall()




    if request.method == "POST":
        butt = html.escape(request.form['butt'])
        if butt == 'DocOrd':
            print('DocOrd')
            Emaillo = html.escape(request.form['Emm'])
            conn = sqlite3.connect('Totringdb.db')
            cursor = conn.cursor()
            Emailll = Emaillo
            print(Emailll)
            cursor.execute(f"SELECT * FROM usersz WHERE Email = '{Emailll}'")

            hisa = cursor.fetchone()
            doc = Document("atte.docx")

            table = doc.tables[1]
            table.alignment = WD_ALIGN_VERTICAL.CENTER

            cursor.execute(f"SELECT i, Day, Date, StartTime, EndTime, Hours, WorkType FROM attendees WHERE Email = '{Emailll}'")
            User_Atten = cursor.fetchall()
            print(User_Atten)

            if not User_Atten:
                # If User_Atten is empty, create an empty table
                num_rows = 1
                num_cols = 7
                table = doc.add_table(num_rows, num_cols)
                for cell in table.rows[0].cells:
                    cell.text = ""

            else:
                num_rows = len(User_Atten)
                num_cols = len(User_Atten[0])

                for row_num in range(num_rows):
                    for col_num in range(num_cols):
                        cell = table.cell(row_num + 1, col_num)
                        cell.text = str(User_Atten[row_num][col_num])
                        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

            replacements = [
                ('studName',     hisa[1]),
                ('studid',       hisa[2][1:11]),
                ('colg',         hisa[8]),
                ('studph',       hisa[4]),
                ('studipan',     hisa[10]),
                ('studpankname', hisa[9]),
            ]

            for old_text, new_text in replacements:
                for paragraph in doc.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

            doc.save(f"{hisa[2][1:11]}_At.docx")

            p = f"{hisa[2][1:11]}_At.docx"

            return jsonify('https://www.uhb-work.com/download_file/' + p)

        if butt == 'AQ':
            Emaillo = html.escape(request.form['Emm'])
            conn = sqlite3.connect('Totringdb.db')
            cursor = conn.cursor()
            Emailll = Emaillo
            print(Emailll)
            cursor.execute(f"SELECT * FROM usersz WHERE Email = '{Emailll}'")

            hisa = cursor.fetchone()
            doc = Document("aqqed.docx")

            table = doc.tables[1]
            table.alignment = WD_ALIGN_VERTICAL.CENTER



            replacements = [
                ('studName',     hisa[1]),
                ('studid',       hisa[2][1:11]),
                ('colg',         hisa[8]),
                ('studph',       hisa[4]),
                ('studipan',     hisa[10]),
                ('studpankname', hisa[9]),
                ('أسم الطالبببببببببببببببببببببببببببببببببببببببببببببببببببب',     hisa[1]),
            ]

            for old_text, new_text in replacements:
                for paragraph in doc.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

            doc.save(f"{hisa[2][1:11]}AQ.docx")
            p = f"{hisa[2][1:11]}AQ.docx"



            return jsonify('https://www.uhb-work.com/download_file/'+p)


        if butt == 'AcceptOrder':
            Email = html.escape(request.form.get('Email'))
            i = html.escape(request.form.get('i'))

            mydb.execute(f'UPDATE Orders SET OrderStatus="True" WHERE Email="{Email}" AND i="{i}"')
            db.commit()
            db.close()
            mzil = Mailer(email="", password="")
            mzil.send(receiver=Email,
                    subject="UHB-WORK",
                    message=" تم قبول طلبك يرجى مراجعة صفحة الطلبات الخاصه بك " )
            return redirect('/adminCCA')
        elif butt == "RejectOrder":

            Email = html.escape(request.form.get('Email'))
            i = html.escape(request.form.get('i'))
            mydb.execute(f'DELETE FROM Orders WHERE Email="{Email}" AND i="{i}"')
            db.commit()
            db.close()
            mzil = Mailer(email="", password="")
            mzil.send(receiver=Email,
                    subject="UHB-WORK",
                    message=" نعتذر لعدم قبول طلبكم , يرجى مراجعة صفحت الطلبات الخاصه بكم" )
            return redirect('/adminCCA')



        if butt == 'Acceptuser':
            Email = html.escape(request.form.get('Email'))
            i = html.escape(request.form.get('i'))
            mydb.execute(f'UPDATE usersz SET AdminAccept="True" WHERE Email="{Email}" AND i="{i}"')
            db.commit()
            db.close()
            mzil = Mailer(email="", password="")
            mzil.send(receiver=Email,
                    subject="UHB-WORK",
                    message=" تم قبول طلبك للعمل الجزئي يمكنك الأن الدخول الى المنصة" )
            return redirect('/adminCCA')
        elif butt == "Rejectuser":
            i = html.escape(request.form.get('i'))
            Email = html.escape(request.form.get('Email'))
            mydb.execute(f'DELETE FROM usersz WHERE Email="{Email}" AND i="{i}"')
            db.commit()
            db.close()
            mzil = Mailer(email="", password="")
            mzil.send(receiver=Email,
                    subject="UHB-WORK",
                    message="نعتذر لعدم قبولك في العمل الجزئي" )
            return redirect('/adminCCA')




    return render_template('admin.html',name=nameee,Email=Emailll,ss=ss, alluser=allUser,  whitingAccept=whitingAccept,  whitingOrder=whitingOrder,  OrderAccepted=OrderAccepted)




@app.route('/myord', methods =["GET", "POST"] )
def homed():
    nameee = session.get('Name')
    Emailll = session.get('Email')
    ss = session.get('Email')
    print(session.get('Email'))
    db = sqlite3.connect("Totringdb.db")
    mydb = db.cursor()
    mydb.execute(f"SELECT * FROM Orders WHERE Email='{session.get('Email')}'")
    orders = mydb.fetchall()
    if request.method == 'POST':
        email = session.get('Email')
        name = session.get('Name')
        phone = session.get('Phone')

        description = html.escape(request.form['description'])
        Title = html.escape(request.form['Title'])
        file_url = html.escape(request.form['Url']) or ''
        date = date_
        order_status = 'False'
        mydb.execute('''INSERT INTO Orders (Email, Name, Phone, description, FileUrl, Date, OrderStatus, Title) VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',(email, name, phone, description, file_url, date, order_status, Title))
        db.commit()
        db.close()
        return redirect('/myord')
    return render_template('myord.html',name=nameee,Email=Emailll,ss=ss, allord=orders)





@app.route('/vcode', methods =["GET", "POST"] )
def homef():
    ss = session.get('Vmail')
    db = sqlite3.connect("Totringdb.db")
    mydb = db.cursor()

    if request.method == 'POST':
        vcode = html.escape(request.form.get('Vcode'))
        if int(vcode) == int(session.get('Vcode')):
            mydb.execute(f'UPDATE usersz SET Status="True" WHERE Email="{html.escape(session.get("Vmail"))}"')
            db.commit()
            db.close()
            session.clear()
            return render_template('login.html',ss=None, Whitngmassge='Witng to Accept you From Admin')
        else:

            return redirect('/vcode')
    mzil = Mailer(email="", password="")
    mzil.send(receiver=session.get('Vmail'),
            subject="THIS IS EMAIL FOR VERIFY",
            message=" your verify code is ("+str(session.get('Vcode') )+")" )
    return render_template('vcode.html',ss=None, Email=session.get('Vmail'))



@app.route('/', methods =["GET", "POST"] )
def loginn():
    ss = session.get('Email')
    db = sqlite3.connect("Totringdb.db")
    mydb = db.cursor()
    if request.method == 'POST':
        user = html.escape(request.form.get('user'))
        passe = html.escape(request.form.get('passs'))

        mydb.execute(f"SELECT * FROM usersz WHERE  Email = '{user}' AND Password = '{encrippt(passe)}'")
        AcountValues = mydb.fetchone()
        if AcountValues == None:
            return render_template('login.html',ss=ss, errormassge='Errore User Or password')
        elif AcountValues[2] == 'admin email':
            session['Email'] = 'admin email'
            session['Name'] = 'admin name'
            return redirect('/adminCCA')

        else:
            if AcountValues[5] == 'False':
                session['Vcode'] = random.randint(100011,900019)
                session['Vmail'] = AcountValues[2]
                return redirect('/vcode')

            elif AcountValues[5] == 'True':
                if AcountValues[7] == 'True':
                    session['Name'] = AcountValues[1]
                    session['Email'] = AcountValues[2]
                    session['Phone'] = AcountValues[4]
                    name = session.get('Name')
                    Email = session.get('Email')
                    return redirect('/home')
                elif AcountValues[7] == 'False':
                    return render_template('login.html',ss=ss, errormassge='Witng to Accept you From Admin')
                redirect('/myord')
    return render_template('login.html',ss=ss)


@app.route('/re', methods =["GET", "POST"])
def signup():
    session.clear()
    db = sqlite3.connect("Totringdb.db")
    mydb = db.cursor()
    ss = session.get('Email')
    if request.method == "POST":
        _Name = html.escape(request.form.get('Name') )
        _Email = html.escape(request.form.get('Email') )
        _password = html.escape(request.form.get('Password') )
        _Phone = html.escape(request.form.get('Phone') )
        _college = html.escape(request.form.get('college') )
        _Bank_Name = html.escape(request.form.get('Bank_Name') )
        _iPAN = html.escape(request.form.get('IPAN') )

        mydb.execute("SELECT Email FROM usersz")
        results = mydb.fetchall()
        if (_Email,) in results:
            return render_template('login.html',ss=ss, errormassge='You Alrrdy Have Acount')
        else:
            mydb.execute(f"INSERT INTO usersz ( Name, Email, Password, Phone, Status,  AdminAccept, college, Bank_Name, IPAN)  VALUES ('{_Name}','{_Email}','{encrippt(_password)}','{_Phone}','{'False'}','{'False'}','{_college}','{_Bank_Name}','{_iPAN}')")
            db.commit()
            db.close()
            session['Vmail'] = _Email
            session['Vcode'] = random.randint(100011,900019)
            print(session.get('Vcode'))
            return redirect('/vcode')

@app.route('/atte', methods =["GET", "POST"])
def Atten():
    conn = sqlite3.connect('Totringdb.db')
    cursor = conn.cursor()
    nameee = session.get('Name')
    Emailll = session.get('Email')
    ss = session.get('Email')
    cursor.execute(f"SELECT * FROM attendees WHERE  Email = '{Emailll}'")
    User_Atten = cursor.fetchall()
    if request.method == "POST":
        day = html.escape(request.form['day'])
        date = html.escape(request.form['date'])
        start_time = html.escape(request.form['start_time'])
        end_time = html.escape(request.form['end_time'])
        hours = html.escape(request.form['hours'])
        work_type = html.escape(request.form['work_type'])
        Datai = html.escape(request.form['Datai'])
        conn = sqlite3.connect('Totringdb.db')
        cursor = conn.cursor()



        cursor.execute("INSERT INTO attendees (Day, Date, StartTime, EndTime, Hours, WorkType, Email, Detai, Namee) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (day, date, start_time, end_time, hours, work_type, Emailll, Datai, nameee))
        conn.commit()

        conn.close()
        return redirect('/atte')



    return render_template('attendance.html',name=nameee,Email=Emailll,ss=ss,User_Atten=User_Atten)



@app.route('/Logout', methods =["GET", "POST"])
def Logout():
    session.clear()
    return redirect('/')




@app.route('/download_file/<filename>')
def download_file(filename):

    p = filename
    return send_file(p, as_attachment=True)

if __name__=='__main__':
   app.run(debug=True)